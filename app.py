from flask import Flask, request, send_file, abort, jsonify, after_this_request, current_app
import os, tempfile, shutil, subprocess, zipfile, logging, re
os.environ["OMP_THREAD_LIMIT"] = "1"
from werkzeug.utils import secure_filename
from flask_cors import CORS

# ======================================================
# BASIC CONFIG
# ======================================================
logging.getLogger("werkzeug").setLevel(logging.ERROR)

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

os.environ.setdefault("UNO_PATH", "/usr/lib/libreoffice/program")
os.environ["PATH"] += ":/usr/lib/libreoffice/program:/usr/bin:/usr/local/bin"
POPPLER_PATH = "/usr/bin"

app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024

PER_TOOL_LIMIT_BYTES = {
    "default": 25 * 1024 * 1024,
    "compress-pdf": 50 * 1024 * 1024,
}

MAX_OCR_PAGES = 30
OCR_DPI = 150
PDF_TO_JPG_DPI = 200
IMAGE_THREAD_COUNT = 1
SUBPROCESS_TIMEOUT = 120

# ======================================================
# LAZY IMPORTS
# ======================================================
def lazy_pdf2docx_converter():
    from pdf2docx import Converter
    return Converter

def lazy_pil_Image():
    from PIL import Image
    return Image

def lazy_pdf2image_convert():
    from pdf2image import convert_from_path
    return convert_from_path

def lazy_pytesseract():
    import pytesseract
    return pytesseract

def lazy_pikepdf():
    import pikepdf
    return pikepdf

def lazy_pypdf():
    from PyPDF2 import PdfReader, PdfWriter, PdfMerger
    return PdfReader, PdfWriter, PdfMerger

def lazy_pdfplumber():
    import pdfplumber
    return pdfplumber

def lazy_docx_Document():
    from docx import Document
    return Document

def with_filename(response, filename):
    response.headers["X-Filename"] = filename
    return response

# ======================================================
# TEMP FILE HELPERS
# ======================================================
def tmp_file(ext=""):
    fd, path = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    return path

def tmp_dir():
    return tempfile.mkdtemp()

def cleanup(path):
    try:
        if not path:
            return
        if os.path.isdir(path):
            shutil.rmtree(path, ignore_errors=True)
        elif os.path.isfile(path):
            os.remove(path)
    except Exception:
        pass

# ======================================================
# SAFE STREAMED UPLOAD (SIZE ENFORCED)
# ======================================================
def get_limit_for_tool(tool):
    return PER_TOOL_LIMIT_BYTES.get(tool, PER_TOOL_LIMIT_BYTES["default"])

def save_upload(file_obj, ext=None, max_bytes=None):
    filename = secure_filename(file_obj.filename or "upload")
    extension = ext if ext else os.path.splitext(filename)[1]
    path = tmp_file(extension)

    total = 0
    with open(path, "wb") as f:
        for chunk in iter(lambda: file_obj.stream.read(65536), b""):
            total += len(chunk)
            if max_bytes and total > max_bytes:
                cleanup(path)
                raise ValueError("File too large")
            f.write(chunk)
    return path

def check_request_size_from_files(files, tool):
    limit = get_limit_for_tool(tool)
    if request.content_length and request.content_length > limit:
        return False, "Upload too large"
    return True, ""

# ======================================================
# SUBPROCESS SAFETY
# ======================================================
def run_subprocess(cmd, timeout=SUBPROCESS_TIMEOUT):
    subprocess.run(
        cmd,
        timeout=timeout,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=True
    )

# ======================================================
# TEXT FORMATTING (UNCHANGED LOGIC)
# ======================================================
BULLET_PATTERNS = [
    r'^\s*[-•]\s+',
    r'^\s*\d+\.\s+',
]

def is_bullet_line(s):
    return any(re.match(p, s.strip()) for p in BULLET_PATTERNS)

def detect_heading(lines):
    heads = set()
    for i, ln in enumerate(lines):
        t = ln.strip()
        if t and 1 <= len(t.split()) <= 8 and t.isupper():
            heads.add(i)
    return heads

def merge_lines_to_paragraphs(raw):
    lines = [l.rstrip() for l in raw.splitlines() if l.strip()]
    heads = detect_heading(lines)
    out = []
    i = 0
    while i < len(lines):
        if i in heads:
            out.append(lines[i])
            i += 1
            continue
        if is_bullet_line(lines[i]):
            b = []
            while i < len(lines) and is_bullet_line(lines[i]):
                b.append(lines[i])
                i += 1
            out.append("\n".join(b))
            continue
        p = [lines[i]]
        i += 1
        while i < len(lines) and i not in heads and not is_bullet_line(lines[i]):
            p.append(lines[i])
            i += 1
        out.append(" ".join(p))
    return "\n\n".join(out)

# ======================================================
# OCR
# ======================================================
def ocr_pdf_to_text(pdf_path, max_pages=MAX_OCR_PAGES, dpi=OCR_DPI):
    convert_from_path = lazy_pdf2image_convert()
    pytesseract = lazy_pytesseract()
    Image = lazy_pil_Image()

    tmpd = tmp_dir()
    try:
        imgs = convert_from_path(
            pdf_path,
            dpi=dpi,
            poppler_path=POPPLER_PATH,
            output_folder=tmpd,
            fmt="png",
            paths_only=True,
            thread_count=1,
        )[:max_pages]

        texts = []
        for i, p in enumerate(imgs, 1):
            with Image.open(p).convert("L") as im:
                txt = pytesseract.image_to_string(im, config="--psm 3")
                texts.append(txt)

        return merge_lines_to_paragraphs("\n".join(texts))
    finally:
        cleanup(tmpd)

# ======================================================
# ROUTES (UNCHANGED NAMES)
# ======================================================
@app.post("/rotate-pdf")
def rotate_pdf():
    tool = "rotate-pdf"
    ok, _ = check_request_size_from_files([request.files.get("file")], tool)
    if not ok:
        abort(413)

    angle = int(request.form.get("angle", 90))
    if angle not in (90, 180, 270):
        abort(400, "Angle must be 90, 180, or 270")

    PdfReader, PdfWriter, _ = lazy_pypdf()
    pdf = save_upload(request.files["file"], ".pdf", get_limit_for_tool(tool))
    out = tmp_file(".pdf")

    @after_this_request
    def _c(r):
        cleanup(pdf)
        cleanup(out)
        return r

    reader = PdfReader(pdf)
    writer = PdfWriter()
    for p in reader.pages:
        p.rotate(angle)
        writer.add_page(p)
    with open(out, "wb") as f:
        writer.write(f)
    original_name = os.path.splitext(f.filename)[0]
    download_name = f"{original_name}_rotated.pdf"

    response = send_file(out, as_attachment=True)
    response.headers["X-Filename"] = download_name
    return response


# ======================================================
# ROOT
# ======================================================
@app.get("/")
def home():
    return "PDF Tools Backend Running"

@app.after_request
def cors(r):
    r.headers["Access-Control-Allow-Origin"] = "*"
    r.headers["Access-Control-Allow-Headers"] = "*"
    r.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return r

@app.before_request
def preflight():
    if request.method == "OPTIONS":
        return app.make_response("")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
# ======================================================
# PDF → WORD (pdf2docx + OCR fallback)
# ======================================================
@app.post("/pdf-to-word")
def pdf_to_word():
    tool = "pdf-to-word"
    f = request.files.get("file")
    if not f:
        abort(400, "No file uploaded")

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    PdfReader, PdfWriter, _ = lazy_pypdf()
    Converter = lazy_pdf2docx_converter()
    Document = lazy_docx_Document()

    pdf_path = save_upload(f, ".pdf", get_limit_for_tool(tool))
    unlocked_pdf = tmp_file(".pdf")
    out_docx = tmp_file(".docx")

    @after_this_request
    def _cleanup(resp):
        cleanup(pdf_path)
        cleanup(unlocked_pdf)
        cleanup(out_docx)
        return resp

    reader = PdfReader(pdf_path, strict=False)
    pdf_to_use = pdf_path

    if getattr(reader, "is_encrypted", False):
        if reader.decrypt("") != 1:
            try:
                pikepdf = lazy_pikepdf()
                with pikepdf.open(pdf_path, password="") as p:
                    p.save(unlocked_pdf)
                pdf_to_use = unlocked_pdf
            except Exception:
                return jsonify({"error": "PDF encrypted. Unlock first."}), 400
        else:
            writer = PdfWriter()
            for p in reader.pages:
                writer.add_page(p)
            with open(unlocked_pdf, "wb") as o:
                writer.write(o)
            pdf_to_use = unlocked_pdf

    try:
        cv = Converter(pdf_to_use)
        cv.convert(out_docx, start=0, end=None)
        cv.close()

        doc = Document(out_docx)
        if doc.paragraphs:
            return send_file(out_docx, as_attachment=True, download_name="output.docx")
    except Exception:
        pass  # fallback to OCR

    text = ocr_pdf_to_text(pdf_to_use)
    doc = Document()
    for block in text.split("\n\n"):
        if block.startswith("•"):
            for ln in block.splitlines():
                doc.add_paragraph(ln.lstrip("• "), style="List Bullet")
        else:
            doc.add_paragraph(block)
    doc.save(out_docx)

    return send_file(out_docx, as_attachment=True, download_name="output.docx")

# ======================================================
# WORD → PDF
# ======================================================
def safe_libreoffice_convert(input_path, out_dir, convert_filter):
    cmd = [
        "libreoffice",
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--nodefault",
        "--norestore",
        "--convert-to", convert_filter,
        "--outdir", out_dir,
        input_path
    ]
    run_subprocess(cmd)


@app.post("/word-to-pdf")
def word_to_pdf():
    tool = "word-to-pdf"
    f = request.files.get("file")
    if not f:
        abort(400)

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".doc", ".docx"):
        abort(400)

    doc_path = save_upload(f, ext, get_limit_for_tool(tool))
    out_dir = tmp_dir()

    @after_this_request
    def _c(r):
        cleanup(doc_path)
        cleanup(out_dir)
        return r

    # STEP 1: DOC/DOCX → ODT
    safe_libreoffice_convert(doc_path, out_dir, "odt")

    base = os.path.splitext(os.path.basename(doc_path))[0]
    odt_path = os.path.join(out_dir, base + ".odt")

    if not os.path.exists(odt_path):
        abort(500, "ODT conversion failed")

    # STEP 2: ODT → PDF (best quality)
    safe_libreoffice_convert(
        odt_path,
        out_dir,
        "pdf:writer_pdf_Export:EmbedStandardFonts=true"
    )

    out_pdf = os.path.join(out_dir, base + ".pdf")
    if not os.path.exists(out_pdf):
        abort(500, "PDF conversion failed")

    return send_file(out_pdf, as_attachment=True, download_name="output.pdf")


# ======================================================
# PPT → PDF
# ======================================================
@app.post("/ppt-to-pdf")
def ppt_to_pdf():
    tool = "ppt-to-pdf"
    f = request.files.get("file")
    if not f:
        abort(400)

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    ppt = save_upload(f, ".pptx", get_limit_for_tool(tool))
    out_dir = tmp_dir()

    @after_this_request
    def _c(r):
        cleanup(ppt)
        cleanup(out_dir)
        return r

    safe_libreoffice_convert(ppt, out_dir, "pdf")

    base_name = os.path.splitext(f.filename)[0]
    out_pdf = os.path.join(out_dir, base_name + ".pdf")

    resp = send_file(out_pdf, as_attachment=True)
    resp.headers["X-Filename"] = f"{base_name}.pdf"
    return resp

# ======================================================
# JPG → PDF (LOW MEMORY)
# ======================================================
@app.post("/jpg-to-pdf")
def jpg_to_pdf():
    tool = "jpg-to-pdf"
    files = request.files.getlist("files")
    if not files:
        abort(400)

    ok, err = check_request_size_from_files(files, tool)
    if not ok:
        abort(413, err)

    Image = lazy_pil_Image()
    images = []
    saved = []
    out_pdf = tmp_file(".pdf")

    @after_this_request
    def _c(r):
        for p in saved:
            cleanup(p)
        cleanup(out_pdf)
        return r

    for f in files:
        p = save_upload(f, None, get_limit_for_tool(tool))
        saved.append(p)
        with Image.open(p).convert("RGB") as im:
            images.append(im.copy())

    images[0].save(
    out_pdf,
    save_all=True,
    append_images=images[1:],
    dpi=(300, 300),
    quality=95,
    subsampling=0
)

    return send_file(out_pdf, as_attachment=True, download_name="output.pdf")

# ======================================================
# PDF → JPG (PAGE AWARE)
# ======================================================
@app.post("/pdf-to-jpg")
def pdf_to_jpg():
    tool = "pdf-to-jpg"
    f = request.files.get("file")
    if not f:
        abort(400)

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    pages = request.form.get("pages")
    convert_from_path = lazy_pdf2image_convert()

    pdf = save_upload(f, ".pdf", get_limit_for_tool(tool))
    out_dir = tmp_dir()
    zip_path = tmp_file(".zip")

    @after_this_request
    def _c(r):
        cleanup(pdf)
        cleanup(out_dir)
        cleanup(zip_path)
        return r

    # ✅ Parse page ranges
    page_numbers = None
    if pages:
        page_numbers = set()
        for part in pages.split(","):
            if "-" in part:
                a, b = map(int, part.split("-"))
                page_numbers.update(range(a, b + 1))
            else:
                page_numbers.add(int(part))

    imgs = convert_from_path(
        pdf,
        dpi=PDF_TO_JPG_DPI,
        poppler_path=POPPLER_PATH,
        output_folder=out_dir,
        fmt="jpeg",
        paths_only=True,
        thread_count=IMAGE_THREAD_COUNT,
    )

    with zipfile.ZipFile(zip_path, "w") as z:
        for i, p in enumerate(imgs, 1):
            if page_numbers and i not in page_numbers:
                continue
            z.write(p, f"page_{i}.jpg")

    orig = os.path.splitext(f.filename)[0]
    resp = send_file(zip_path, as_attachment=True)
    return with_filename(resp, f"{orig}_images.zip")


# ======================================================
# MERGE PDF
# ======================================================
@app.post("/merge-pdf")
def merge_pdf():
    tool = "merge-pdf"
    files = request.files.getlist("files")
    ok, err = check_request_size_from_files(files, tool)
    if not ok:
        abort(413, err)

    _, _, PdfMerger = lazy_pypdf()
    merger = PdfMerger()
    out = tmp_file(".pdf")

    @after_this_request
    def _c(r):
        cleanup(out)
        return r

    for f in files:
        p = save_upload(f, ".pdf", get_limit_for_tool(tool))
        merger.append(p)
    merger.write(out)
    merger.close()
    resp = send_file(out, as_attachment=True)
    return with_filename(resp, "merged.pdf")


# ======================================================
# SPLIT PDF
# ======================================================
@app.post("/split-pdf")
def split_pdf():
    tool = "split-pdf"
    f = request.files.get("file")
    ranges = request.form.get("ranges")
    if not f or not ranges:
        abort(400)

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    PdfReader, PdfWriter, _ = lazy_pypdf()
    pdf = save_upload(f, ".pdf", get_limit_for_tool(tool))
    out_dir = tmp_dir()
    zip_path = tmp_file(".zip")

    @after_this_request
    def _c(r):
        cleanup(pdf)
        cleanup(out_dir)
        cleanup(zip_path)
        return r

    reader = PdfReader(pdf)
    with zipfile.ZipFile(zip_path, "w") as z:
        for r in ranges.split(","):
            a, b = map(int, r.split("-")) if "-" in r else (int(r), int(r))
            for i in range(a, b + 1):
                w = PdfWriter()
                w.add_page(reader.pages[i - 1])
                p = os.path.join(out_dir, f"page_{i}.pdf")
                with open(p, "wb") as o:
                    w.write(o)
                z.write(p, f"page_{i}.pdf")

    return send_file(zip_path, as_attachment=True, download_name="split.zip")

# ======================================================
# COMPRESS PDF
# ======================================================
@app.post("/compress-pdf")
def compress_pdf():
    tool = "compress-pdf"
    f = request.files.get("file")
    if not f:
        abort(400)

    if not shutil.which("gs"):
        abort(500, "Ghostscript not installed")

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    level = request.form.get("level", "screen")

    GS_LEVELS = {
        "low": "/screen",     # maximum compression
        "medium": "/ebook",   # balanced
        "high": "/printer",   # high quality
    }

    inp = save_upload(f, ".pdf", get_limit_for_tool(tool))
    out = tmp_file(".pdf")

    @after_this_request
    def _c(r):
        cleanup(inp)
        cleanup(out)
        return r

    cmd = [
        "gs",
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS={GS_LEVELS.get(level, '/screen')}",
        "-dNOPAUSE",
        "-dBATCH",
        "-dQUIET",
        "-dDetectDuplicateImages=true",
        "-dCompressFonts=true",
        "-dDownsampleColorImages=true",
        "-dColorImageResolution=150",
        f"-sOutputFile={out}",
        inp
    ]

    run_subprocess(cmd)
    orig = os.path.splitext(f.filename)[0]
    final_name = f"{orig}_compressed.pdf"

    resp = send_file(out, as_attachment=True)
    return with_filename(resp, final_name)


# ======================================================
# PROTECT / UNLOCK PDF
# ======================================================
@app.post("/protect-pdf")
def protect_pdf():
    tool = "protect-pdf"
    f = request.files.get("file")
    pwd = request.form.get("password")
    if not f or not pwd:
        abort(400)

    PdfReader, PdfWriter, _ = lazy_pypdf()
    pdf = save_upload(f, ".pdf", get_limit_for_tool(tool))
    out = tmp_file(".pdf")

    @after_this_request
    def _c(r):
        cleanup(pdf)
        cleanup(out)
        return r

    r = PdfReader(pdf)
    w = PdfWriter()
    for p in r.pages:
        w.add_page(p)
    w.encrypt(pwd)
    with open(out, "wb") as o:
        w.write(o)

    return send_file(out, as_attachment=True, download_name="protected.pdf")

@app.post("/unlock-pdf")
def unlock_pdf():
    tool = "unlock-pdf"
    f = request.files.get("file")
    pwd = request.form.get("password", "")

    if not f:
        abort(400)

    pdf = save_upload(f, ".pdf", get_limit_for_tool(tool))
    out = tmp_file(".pdf")

    @after_this_request
    def _c(r):
        cleanup(pdf)
        cleanup(out)
        return r

    try:
        pikepdf = lazy_pikepdf()
        with pikepdf.open(pdf, password=pwd) as p:
            p.save(out)
    except Exception:
        abort(400, "Wrong password")

    return send_file(out, as_attachment=True, download_name="unlocked.pdf")

# ======================================================
# EXTRACT TEXT
# ======================================================
@app.post("/extract-text")
def extract_text():
    tool = "extract-text"
    f = request.files.get("file")
    if not f:
        abort(400)

    ok, err = check_request_size_from_files([f], tool)
    if not ok:
        abort(413, err)

    pdf = save_upload(f, ".pdf", get_limit_for_tool(tool))
    pdfplumber = lazy_pdfplumber()

    @after_this_request
    def _c(r):
        cleanup(pdf)
        return r

    text = ""
    with pdfplumber.open(pdf) as p:
        for pg in p.pages:
            t = pg.extract_text()
            if t:
                text += t + "\n"

    if not text.strip():
        text = ocr_pdf_to_text(pdf)

    original_name = os.path.splitext(f.filename)[0]

    return jsonify({
        "text": merge_lines_to_paragraphs(text),
        "filename": f"{original_name}.txt"
    })